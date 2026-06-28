# app/services/parser.py
# Full path: /Users/vantakujagadeesh/Desktop/rag n8n/autorag/app/services/parser.py
#
# Document parser supporting PDF, DOCX, and URL sources.
# Each parser extracts raw text and returns a cleaned string.
# ──────────────────────────────────────────────────────────────────────────────

import io
import re
import logging
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from playwright.async_api import async_playwright, TimeoutError as PlaywrightTimeout

logger = logging.getLogger(__name__)

# File extensions we support
SUPPORTED_EXTENSIONS: dict[str, str] = {
    ".pdf": "pdf",
    ".docx": "docx",
}


# ── Helpers ───────────────────────────────────────────────────────────────────

def _clean_text(raw: str) -> str:
    """
    Normalise extracted text:
      - collapse multiple blank lines into one
      - strip leading/trailing whitespace per line
      - remove null bytes and control chars (except newline/tab)
    """
    # Remove null bytes and non-printable control characters (keep \n, \t)
    cleaned = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", raw)
    # Collapse 3+ consecutive newlines into 2
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    # Strip trailing whitespace on each line
    cleaned = "\n".join(line.rstrip() for line in cleaned.splitlines())
    return cleaned.strip()


# ── PDF Parser ────────────────────────────────────────────────────────────────

def parse_pdf(file_bytes: bytes) -> str:
    """
    Extract text from a PDF using PyMuPDF (fitz).

    Args:
        file_bytes: Raw PDF file content.

    Returns:
        Cleaned concatenated text from all pages.

    Raises:
        ValueError: If the PDF is encrypted and cannot be opened,
                    or if zero text could be extracted.
    """
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as e:
        logger.warning(f"Failed to open PDF: {e}")
        raise ValueError(f"Cannot open PDF file: {e}") from e

    if doc.is_encrypted:
        try:
            # Attempt empty-password unlock (some PDFs set owner-only restrictions)
            if not doc.authenticate(""):
                logger.warning("PDF is encrypted and could not be unlocked")
                raise ValueError(
                    "PDF is encrypted. Please provide an unencrypted version."
                )
        except Exception as e:
            logger.warning(f"PDF encryption check failed: {e}")
            raise ValueError(f"PDF encryption error: {e}") from e

    pages: list[str] = []
    for page_num in range(len(doc)):
        try:
            page = doc[page_num]
            text = page.get_text("text")
            if text and text.strip():
                pages.append(text)
        except Exception as e:
            logger.warning(
                f"Failed to extract text from PDF page {page_num + 1}: {e}"
            )
            # Continue with remaining pages rather than aborting
            continue

    page_count = len(doc)
    doc.close()

    if not pages:
        logger.warning("PDF produced zero extractable text (possibly scanned/image-only)")
        raise ValueError(
            "No text could be extracted from this PDF. "
            "It may be a scanned document requiring OCR."
        )

    raw = "\n\n".join(pages)
    cleaned = _clean_text(raw)
    logger.info(f"PDF parsed: {page_count} pages → {len(cleaned):,} chars")
    return cleaned


# ── DOCX Parser ───────────────────────────────────────────────────────────────

def parse_docx(file_bytes: bytes) -> str:
    """
    Extract text from a DOCX file using python-docx.
    Extracts both paragraph text and table cell text.

    Args:
        file_bytes: Raw DOCX file content.

    Returns:
        Cleaned text with paragraphs joined by newlines.

    Raises:
        ValueError: If the file cannot be parsed as a valid DOCX.
    """
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
    except Exception as e:
        logger.warning(f"Failed to open DOCX: {e}")
        raise ValueError(f"Cannot open DOCX file: {e}") from e

    parts: list[str] = []

    # Extract paragraphs (skip empty ones)
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Extract table contents row by row
    for table in doc.tables:
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_cells.append(cell_text)
            if row_cells:
                parts.append(" | ".join(row_cells))

    if not parts:
        logger.warning("DOCX produced zero extractable text")
        raise ValueError("No text could be extracted from this DOCX file.")

    raw = "\n".join(parts)
    cleaned = _clean_text(raw)
    logger.info(
        f"DOCX parsed: {len(doc.paragraphs)} paragraphs, "
        f"{len(doc.tables)} tables → {len(cleaned):,} chars"
    )
    return cleaned


# ── URL Parser ────────────────────────────────────────────────────────────────

async def parse_url(url: str) -> str:
    """
    Fetch a web page via headless Chromium (Playwright) and extract body text.

    - Waits for networkidle to ensure SPAs finish rendering.
    - 30-second timeout.
    - Browser is always closed after extraction, even on failure.

    Args:
        url: Fully qualified URL (https://...).

    Returns:
        Cleaned innerText of the page body.

    Raises:
        ValueError: If the page fails to load or produces no text.
    """
    logger.info(f"Fetching URL via Playwright: {url}")
    playwright = None
    browser = None

    try:
        playwright = await async_playwright().start()
        browser = await playwright.chromium.launch(headless=True)

        context = await browser.new_context(
            user_agent=(
                "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/125.0.0.0 Safari/537.36"
            ),
            java_script_enabled=True,
        )

        page = await context.new_page()

        try:
            await page.goto(url, wait_until="networkidle", timeout=30_000)
        except PlaywrightTimeout:
            logger.warning(f"URL timed out after 30s: {url}")
            raise ValueError(f"Page load timed out after 30 seconds: {url}")

        # Extract visible text from the body element
        body = await page.query_selector("body")
        if body is None:
            raise ValueError(f"No <body> element found on page: {url}")

        raw_text = await body.inner_text()

        if not raw_text or not raw_text.strip():
            logger.warning(f"URL produced zero text: {url}")
            raise ValueError(f"No text could be extracted from: {url}")

        cleaned = _clean_text(raw_text)
        logger.info(f"URL parsed: {url} → {len(cleaned):,} chars")
        return cleaned

    except ValueError:
        # Re-raise our own ValueErrors as-is
        raise
    except Exception as e:
        logger.error(f"Unexpected error fetching URL {url}: {e}")
        raise ValueError(f"Failed to fetch and parse URL: {e}") from e
    finally:
        # Always clean up browser resources
        if browser is not None:
            try:
                await browser.close()
            except Exception:
                pass
        if playwright is not None:
            try:
                await playwright.stop()
            except Exception:
                pass


# ── File Type Detection ──────────────────────────────────────────────────────

def detect_file_type(filename: str) -> str:
    """
    Detect document type from filename extension.

    Args:
        filename: Original filename (e.g. "report.pdf").

    Returns:
        "pdf" or "docx".

    Raises:
        ValueError: If the file extension is not supported.
    """
    ext = Path(filename).suffix.lower()

    if ext in SUPPORTED_EXTENSIONS:
        return SUPPORTED_EXTENSIONS[ext]

    supported_list = ", ".join(sorted(SUPPORTED_EXTENSIONS.keys()))
    raise ValueError(
        f"Unsupported file type '{ext}' for file '{filename}'. "
        f"Supported extensions: {supported_list}"
    )


# ── Unified File Parser ─────────────────────────────────────────────────────

def parse_file(filename: str, file_bytes: bytes) -> str:
    """
    Auto-detect file type from filename and parse accordingly.

    Args:
        filename:   Original filename (e.g. "policy.pdf").
        file_bytes: Raw file content.

    Returns:
        Cleaned extracted text.

    Raises:
        ValueError: If file type is unsupported or parsing fails.
    """
    file_type = detect_file_type(filename)
    logger.info(f"Parsing file: '{filename}' (detected type: {file_type})")

    if file_type == "pdf":
        text = parse_pdf(file_bytes)
    elif file_type == "docx":
        text = parse_docx(file_bytes)
    else:
        raise ValueError(f"No parser available for type: {file_type}")

    logger.info(
        f"Parse complete: '{filename}' → {len(text):,} chars extracted"
    )
    return text
